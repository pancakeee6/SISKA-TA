from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Query, Form, Request
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, func, case
from sqlalchemy.orm import selectinload
from datetime import datetime, timezone, timedelta
import logging
import os
import uuid
import io
import json

from app.db.database import get_db
from app.api.deps import get_current_admin
from app.models.attendance import AttendanceLog
from app.models.user import User
from app.models.shift import WorkShift
from app.schemas.attendance import AttendanceLogResponse, AttendanceListResponse, AttendanceDinasCreate
from app.services import ai_service
from app.core.websocket import ws_manager
from app.core.config import settings
from fastapi.responses import StreamingResponse

logger = logging.getLogger(__name__)

router = APIRouter()

def count_working_days(start_date, end_date):
    if not start_date or not end_date:
        return 0
    days = 0
    curr = start_date
    while curr <= end_date:
        if curr.weekday() < 6: # 0=Mon, 5=Sat, 6=Sun
            days += 1
        curr += timedelta(days=1)
    return days


# Threshold untuk filter confidence/distance dari AI
# Jika AI mengembalikan confidence < threshold, dianggap tidak yakin
MIN_CONFIDENCE = 0.6
MAX_DISTANCE = 0.7


@router.post("/recognize")
async def recognize_attendance(
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
):
    """Receive face image, forward to AI API, log attendance."""
    try:
        # Forward to AI API
        ai_result = await ai_service.recognize_face(file)
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"AI API error: {str(e)}")

    # ─── LOG RAW AI RESPONSE untuk debugging salah deteksi ───
    logger.info(f"[AI RAW RESPONSE] {ai_result}")

    if not ai_result.get("faces") or ai_result.get("status") in ("error", "failed", "unrecognized"):
        return {
            "status": "unrecognized",
            "message": "Wajah tidak dikenali",
            "faces": [],
        }

    # Process each recognized face
    results = []
    for face in ai_result["faces"]:
        # ─── Filter berdasarkan confidence/distance jika AI mengembalikannya ───
        confidence = face.get("confidence") or face.get("score") or face.get("similarity")
        distance = face.get("distance")
        
        if confidence is not None and float(confidence) < MIN_CONFIDENCE:
            logger.warning(f"[SKIP LOW CONFIDENCE] {face.get('name')} confidence={confidence} < {MIN_CONFIDENCE}")
            continue
        if distance is not None and float(distance) > MAX_DISTANCE:
            logger.warning(f"[SKIP HIGH DISTANCE] {face.get('name')} distance={distance} > {MAX_DISTANCE}")
            continue
        
        # Find user by ml_person_id first
        ml_person_id = face.get("id") or face.get("person_id")
        user = None
        
        if ml_person_id is not None:
            user_result = await db.execute(
                select(User).where(
                    User.ml_person_id == ml_person_id,
                    User.is_active == True
                )
            )
            user = user_result.scalars().first()
            
        # Fallback to name if ml_person_id not found or failed
        if not user:
            user_result = await db.execute(
                select(User).where(
                    func.lower(User.full_name) == func.lower(face.get("name", "")),
                    User.is_active == True
                )
            )
            user = user_result.scalars().first()

        if user:
            # ─── Evaluasi Shift, Waktu WIB, Toleransi 20 Menit, & Mencegah Duplikasi ───
            now_utc = datetime.now(timezone.utc)
            wib_tz = timezone(timedelta(hours=7))
            now_wib = now_utc.astimezone(wib_tz)
            today_date = now_wib.date()

            # Ambil semua WorkShift dari database atau gunakan default
            shift_result = await db.execute(select(WorkShift).order_by(WorkShift.start_time.asc()))
            shifts = shift_result.scalars().all()
            if not shifts:
                shifts = [
                    WorkShift(name="Shift Pagi", start_time="08:00", end_time="15:00"),
                    WorkShift(name="Shift Sore", start_time="15:00", end_time="21:00")
                ]

            # Tentukan shift mana yang aktif berdasarkan jam saat ini (WIB)
            active_shift = shifts[0]
            for s in shifts:
                try:
                    s_hour, s_min = map(int, s.start_time.split(":"))
                    if now_wib.time() >= datetime(2000, 1, 1, s_hour, s_min).time():
                        active_shift = s
                except Exception:
                    pass

            # Cari log kehadiran user hari ini (dalam rentang waktu WIB hari ini)
            start_of_day_wib = datetime.combine(today_date, datetime.min.time(), tzinfo=wib_tz)
            start_of_day_utc = start_of_day_wib.astimezone(timezone.utc)
            end_of_day_utc = start_of_day_utc + timedelta(days=1)

            existing_logs_result = await db.execute(
                select(AttendanceLog)
                .where(
                    AttendanceLog.user_id == user.id,
                    AttendanceLog.timestamp >= start_of_day_utc,
                    AttendanceLog.timestamp < end_of_day_utc
                )
                .order_by(AttendanceLog.timestamp.asc())
            )
            today_logs = existing_logs_result.scalars().all()

            in_logs = [l for l in today_logs if l.event_type == "IN"]
            out_logs = [l for l in today_logs if l.event_type == "OUT"]

            # Jika sudah menyelesaikan 1 shift hari ini (1 IN, 1 OUT lengkap) dan mencoba IN lagi, 
            # pastikan shiftnya berlanjut ke shift ke-2 (Sore) tanpa peduli jam awal.
            if len(in_logs) == len(out_logs) and len(in_logs) == 1:
                if len(shifts) > 1:
                    active_shift = shifts[1]

            skip_log = False
            calculated_event_type = "IN"
            calculated_is_late = False
            status_text = "present"
            late_duration_str = None
            diff_minutes_late = None

            if len(in_logs) > len(out_logs):
                # Sedang berada di dalam shift, mencoba OUT
                unclosed_in_time = in_logs[-1].timestamp
                if unclosed_in_time.tzinfo is None:
                    unclosed_in_time = unclosed_in_time.replace(tzinfo=timezone.utc)
                diff_minutes = (now_utc - unclosed_in_time).total_seconds() / 60.0
                
                calculated_event_type = "OUT"
                calculated_is_late = False

                if diff_minutes < 15.0:
                    # Cegah duplikasi spam dalam 15 menit setelah IN
                    skip_log = True
                    status_text = "cooldown"
                else:
                    # Validasi Time Window untuk OUT
                    # Cari shift dari waktu IN awal untuk menentukan batas pulang (end_time) yang benar
                    unclosed_wib = unclosed_in_time.astimezone(wib_tz)
                    original_shift = shifts[0]
                    for s in shifts:
                        try:
                            s_hour, s_min = map(int, s.start_time.split(":"))
                            if unclosed_wib.time() >= datetime(2000, 1, 1, s_hour, s_min).time():
                                original_shift = s
                        except Exception:
                            pass
                            
                    # Terapkan label shift berdasarkan shift awal (saat IN)
                    active_shift = original_shift
                            
                    # Gunakan original_shift.end_time secara mutlak sebagai batas minimal waktu Pulang
                    now_time = now_wib.time()
                    try:
                        end_h, end_m = map(int, original_shift.end_time.split(":"))
                        shift_end_time = datetime(2000, 1, 1, end_h, end_m).time()
                    except Exception:
                        shift_end_time = datetime(2000, 1, 1, 12, 0).time()
                    
                    if now_time < shift_end_time:
                        skip_log = True
                        status_text = "early_out"
                    else:
                        status_text = "present"
            else:
                # Sedang di luar shift, mencoba IN
                if len(in_logs) >= 2:
                    # Sudah memenuhi kuota 2 shift per hari
                    skip_log = True
                    calculated_event_type = "IN"
                    calculated_is_late = False
                    status_text = "max_shifts"
                else:
                    calculated_event_type = "IN"
                    if out_logs:
                        last_out = out_logs[-1].timestamp
                        if last_out.tzinfo is None:
                            last_out = last_out.replace(tzinfo=timezone.utc)
                        if (now_utc - last_out).total_seconds() / 60.0 < 15.0:
                            skip_log = True
                            status_text = "cooldown"
                            calculated_is_late = False
                    
                    if not skip_log:
                        # Logika keterlambatan IN
                        try:
                            sh, sm = map(int, active_shift.start_time.split(":"))
                            shift_start_dt = now_wib.replace(hour=sh, minute=sm, second=0, microsecond=0)
                            late_tol = getattr(active_shift, 'late_tolerance', 15)
                            late_threshold_dt = shift_start_dt + timedelta(minutes=late_tol)
                            if now_wib > late_threshold_dt:
                                calculated_is_late = True
                                status_text = "late"
                                diff_minutes_late = int((now_wib - shift_start_dt).total_seconds() / 60.0)
                                diff_minutes_late = max(0, diff_minutes_late)
                                hours = diff_minutes_late // 60
                                mins = diff_minutes_late % 60
                                if hours > 0 and mins > 0:
                                    late_duration_str = f"{hours} jam {mins} menit"
                                elif hours > 0:
                                    late_duration_str = f"{hours} jam"
                                else:
                                    late_duration_str = f"{mins} menit" if mins > 0 else "1 menit"
                            else:
                                calculated_is_late = False
                                status_text = "present"
                        except Exception:
                            calculated_is_late = bool(face.get("is_late", False))
                            status_text = "late" if calculated_is_late else "present"

            if not skip_log:
                log = AttendanceLog(
                    user_id=user.id,
                    event_type=calculated_event_type,
                    status=status_text,
                    late=calculated_is_late,
                    device_id=settings.DEVICE_ID,
                )
                db.add(log)
                await db.commit()

                event_data = {
                    "type": "attendance_marked",
                    "data": {
                        "user_name": user.full_name,
                        "event_type": calculated_event_type,
                        "status": status_text,
                        "late": calculated_is_late,
                        "timestamp": now_utc.isoformat(),
                        "shift_label": active_shift.name,
                        "late_duration": late_duration_str,
                        "late_minutes": diff_minutes_late,
                    },
                }
                await ws_manager.broadcast(event_data)
            
            results.append({
                "user_name": user.full_name,
                "event_type": calculated_event_type,
                "status": "cooldown" if skip_log else "ok",
                "late": calculated_is_late,
                "audio_text": face.get("audio_text"),
                "bbox": face.get("bbox") or face.get("box") or face.get("location"),
                "shift_label": active_shift.name,
                "late_duration": late_duration_str,
                "late_minutes": diff_minutes_late,
            })

    return {
        "status": "recognized",
        "faces": results,
    }


@router.post("/dinas")
async def record_dinas(
    user_id: str = Form(...),
    date: str = Form(None),
    keterangan: str = Form("Perizinan"),
    shift: str = Form("Seharian"),
    file: UploadFile = File(None),
    db: AsyncSession = Depends(get_db),
    _admin=Depends(get_current_admin),
):
    """Record out-of-town official duty or permit (Perizinan) for a user."""
    user_result = await db.execute(select(User).where(User.id == user_id))
    user = user_result.scalars().first()
    if not user:
        raise HTTPException(status_code=404, detail="Pegawai / Dosen tidak ditemukan")

    wib_tz = timezone(timedelta(hours=7))
    if date:
        try:
            target_date = datetime.strptime(date, "%Y-%m-%d").date()
            dt_wib = datetime.combine(target_date, datetime.strptime("08:00:00", "%H:%M:%S").time(), tzinfo=wib_tz)
            target_timestamp = dt_wib.astimezone(timezone.utc)
        except ValueError:
            raise HTTPException(status_code=400, detail="Format tanggal harus YYYY-MM-DD")
    else:
        target_timestamp = datetime.now(timezone.utc)

    # Handle file upload
    attachment_path = None
    if file and file.filename:
        upload_dir = os.path.join(settings.UPLOAD_DIR, "permits")
        os.makedirs(upload_dir, exist_ok=True)
        
        ext = os.path.splitext(file.filename)[1]
        unique_filename = f"{uuid.uuid4().hex}{ext}"
        file_path = os.path.join(upload_dir, unique_filename)
        
        content = await file.read()
        with open(file_path, 'wb') as out_file:
            out_file.write(content)
            
        attachment_path = f"/api/v1/uploads/permits/{unique_filename}"

    import json
    dinas_data = {
        "k": keterangan[:50] if keterangan else "Perizinan",
        "s": shift,
        "f": unique_filename if attachment_path else None
    }
    device_id_val = json.dumps(dinas_data)

    log = AttendanceLog(
        user_id=user.id,
        event_type="DINAS",
        status="dinas",
        late=False,
        device_id=device_id_val,
        timestamp=target_timestamp,
    )
    db.add(log)
    await db.commit()
    await db.refresh(log)

    return {
        "status": "success",
        "message": f"Izin berhasil dicatat untuk {user.full_name}",
        "log_id": log.id
    }


@router.get("/logs", response_model=AttendanceListResponse)
async def attendance_logs(
    page: int = Query(1, ge=1),
    per_page: int = Query(20, ge=1, le=100),
    date: str = Query(None, description="Filter by date (YYYY-MM-DD)"),
    date_from: str = Query(None),
    date_to: str = Query(None),
    status: str = Query(None),
    search: str = Query(None),
    shift: str = Query(None, description="Filter by Shift (Pagi/Sore/Seharian)"),
    db: AsyncSession = Depends(get_db),
    _admin=Depends(get_current_admin),
):
    """Get attendance logs with pagination and filters."""
    query = select(AttendanceLog).join(User)

    wib_tz = timezone(timedelta(hours=7))

    # Date filter
    if date:
        try:
            filter_date = datetime.strptime(date, "%Y-%m-%d").date()
            dt_start = datetime.combine(filter_date, datetime.min.time(), tzinfo=wib_tz).astimezone(timezone.utc)
            dt_end = datetime.combine(filter_date, datetime.max.time(), tzinfo=wib_tz).astimezone(timezone.utc)
            query = query.where(AttendanceLog.timestamp >= dt_start, AttendanceLog.timestamp <= dt_end)
        except ValueError:
            raise HTTPException(status_code=400, detail="Format tanggal harus YYYY-MM-DD")

    # Date range filters
    if date_from:
        try:
            d_from = datetime.strptime(date_from, "%Y-%m-%d").date()
            dt_from = datetime.combine(d_from, datetime.min.time(), tzinfo=wib_tz).astimezone(timezone.utc)
            query = query.where(AttendanceLog.timestamp >= dt_from)
        except ValueError:
            raise HTTPException(status_code=400, detail="Format date_from harus YYYY-MM-DD")

    if date_to:
        try:
            d_to = datetime.strptime(date_to, "%Y-%m-%d").date()
            dt_to = datetime.combine(d_to, datetime.max.time(), tzinfo=wib_tz).astimezone(timezone.utc)
            query = query.where(AttendanceLog.timestamp <= dt_to)
        except ValueError:
            raise HTTPException(status_code=400, detail="Format date_to harus YYYY-MM-DD")

    # Status filter
    if status == "late":
        query = query.where(AttendanceLog.late == True)
    elif status == "present":
        query = query.where(AttendanceLog.late == False, AttendanceLog.status != "dinas", AttendanceLog.event_type != "DINAS")
    elif status == "dinas":
        query = query.where((AttendanceLog.status == "dinas") | (AttendanceLog.event_type == "DINAS"))

    # Search by user name or employee_id
    if search:
        query = query.where(User.full_name.ilike(f"%{search}%") | User.employee_id.ilike(f"%{search}%"))
        
    # Note: Shift filter relies on python-side filtering because shift is determined dynamically, 
    # but we can filter DINAS based on shift JSON, or time for regular attendance.
    # To keep it robust with pagination, we filter heuristically in SQL if possible, 
    # but exact shift is evaluated later. For now, we'll fetch more or apply heuristic.
    if shift and shift.lower() != "all":
        if shift.lower() == "pagi":
            # Jam Masuk < 15:00 UTC+7 (08:00 UTC)
            query = query.where(func.extract('hour', func.timezone('UTC', AttendanceLog.timestamp)) < 8)
        elif shift.lower() == "sore":
            # Jam Masuk >= 15:00 UTC+7 (08:00 UTC)
            query = query.where(func.extract('hour', func.timezone('UTC', AttendanceLog.timestamp)) >= 8)
        elif shift.lower() == "seharian":
            query = query.where(AttendanceLog.event_type == "DINAS")

    # Count
    count_query = select(func.count()).select_from(query.subquery())
    total = (await db.execute(count_query)).scalar() or 0

    # Paginate
    query = query.offset((page - 1) * per_page).limit(per_page).order_by(AttendanceLog.created_at.desc())
    result = await db.execute(query)
    logs = result.scalars().all()

    # Batch fetch all users in 1 query to prevent N+1 queries
    user_ids = {log.user_id for log in logs if log.user_id}
    user_map = {}
    if user_ids:
        users_result = await db.execute(
            select(User).where(User.id.in_(user_ids)).options(selectinload(User.face_data))
        )
        for u in users_result.scalars().all():
            user_map[u.id] = (u.full_name, u.employee_id, u.avatar)

    shift_result = await db.execute(select(WorkShift).order_by(WorkShift.start_time.asc()))
    shifts = shift_result.scalars().all()
    if not shifts:
        shifts = [
            WorkShift(name="Shift 1 (Pagi/Siang)", start_time="08:00", end_time="15:00"),
            WorkShift(name="Shift 2 (Sore/Malam)", start_time="15:00", end_time="21:00")
        ]
    wib_tz = timezone(timedelta(hours=7))

    log_responses = []
    for log in logs:
        user_name, employee_id, avatar = user_map.get(log.user_id, ("Unknown", "-", None))
        resp = AttendanceLogResponse.model_validate(log)
        resp.user_name = user_name
        resp.employee_id = employee_id
        resp.avatar = avatar

        if log.timestamp:
            log_utc = log.timestamp if log.timestamp.tzinfo else log.timestamp.replace(tzinfo=timezone.utc)
            log_wib = log_utc.astimezone(wib_tz)
            active_shift = shifts[0]
            for s in shifts:
                try:
                    sh, sm = map(int, s.start_time.split(":"))
                    if log_wib.time() >= datetime(2000, 1, 1, sh, sm).time():
                        active_shift = s
                except Exception:
                    pass
            if log.status == "dinas" or log.event_type == "DINAS":
                import json
                try:
                    dinas_data = json.loads(log.device_id)
                    resp.shift_label = dinas_data.get("s", "Seharian")
                    resp.device_id = dinas_data.get("k", "Perizinan")
                    if dinas_data.get("f"):
                        resp.attachment_path = f"/api/v1/uploads/permits/{dinas_data['f']}"
                except Exception:
                    resp.shift_label = "Seharian"
                    resp.device_id = log.device_id or "Perizinan"
                
                resp.late_duration = "-"
            else:
                resp.shift_label = active_shift.name
                if log.late and log.event_type == "IN":
                    try:
                        sh, sm = map(int, active_shift.start_time.split(":"))
                        shift_start_dt = log_wib.replace(hour=sh, minute=sm, second=0, microsecond=0)
                        diff_minutes = int((log_wib - shift_start_dt).total_seconds() / 60.0)
                        if diff_minutes < 0:
                            diff_minutes = 0
                        resp.late_minutes = diff_minutes
                        hours = diff_minutes // 60
                        mins = diff_minutes % 60
                        if hours > 0 and mins > 0:
                            resp.late_duration = f"{hours} jam {mins} menit"
                        elif hours > 0:
                            resp.late_duration = f"{hours} jam"
                        else:
                            resp.late_duration = f"{mins} menit" if mins > 0 else "1 menit"
                    except Exception:
                        resp.late_duration = "-"

        log_responses.append(resp)

    return AttendanceListResponse(
        logs=log_responses,
        total=total,
        page=page,
        per_page=per_page,
    )


@router.get("/export")
async def export_attendance(
    request: Request,
    date_from: str = Query(None, description="Start date (YYYY-MM-DD)"),
    date_to: str = Query(None, description="End date (YYYY-MM-DD)"),
    db: AsyncSession = Depends(get_db),
    _admin=Depends(get_current_admin),
):
    """Export attendance data as CSV file."""
    import io
    from fastapi.responses import StreamingResponse
    import openpyxl
    from openpyxl.styles import PatternFill, Font, Alignment, Border, Side

    query = select(AttendanceLog).join(User).order_by(AttendanceLog.timestamp.desc())

    wib_tz = timezone(timedelta(hours=7))

    # Date range filter
    if date_from:
        try:
            d_from = datetime.strptime(date_from, "%Y-%m-%d").date()
            dt_from = datetime.combine(d_from, datetime.min.time(), tzinfo=wib_tz).astimezone(timezone.utc)
            query = query.where(AttendanceLog.timestamp >= dt_from)
        except ValueError:
            raise HTTPException(status_code=400, detail="Format date_from harus YYYY-MM-DD")

    if date_to:
        try:
            d_to = datetime.strptime(date_to, "%Y-%m-%d").date()
            dt_to = datetime.combine(d_to, datetime.max.time(), tzinfo=wib_tz).astimezone(timezone.utc)
            query = query.where(AttendanceLog.timestamp <= dt_to)
        except ValueError:
            raise HTTPException(status_code=400, detail="Format date_to harus YYYY-MM-DD")

    result = await db.execute(query)
    logs = result.scalars().all()

    # Batch fetch users in 1 query
    user_ids = {log.user_id for log in logs if log.user_id}
    user_map = {}
    if user_ids:
        users_result = await db.execute(
            select(User.id, User.full_name, User.employee_id).where(User.id.in_(user_ids))
        )
        for u_id, u_name, u_emp in users_result.all():
            user_map[u_id] = (u_name, u_emp)

    wib_tz = timezone(timedelta(hours=7))
    grouped_logs = {}

    for log in logs:
        if not log.timestamp:
            continue
        dt_wib = log.timestamp.replace(tzinfo=timezone.utc).astimezone(wib_tz)
        date_str = dt_wib.strftime("%Y-%m-%d")
        hour = dt_wib.hour
        shift_label = "Shift Pagi" if hour < 15 else "Shift Sore"
        
        # Gabungkan berdasarkan user dan tanggal saja agar IN dan OUT yang beda shift tetap 1 baris
        key = f"{log.user_id}_{date_str}"
        if key not in grouped_logs:
            grouped_logs[key] = {
                "user_id": log.user_id,
                "date": date_str,
                "shift": shift_label,
                "in_time": "-",
                "out_time": "-",
                "status": "Tidak Hadir",
                "keterangan": "-",
                "device_id": "-",
                "lampiran": "-"
            }
            
        group = grouped_logs[key]
        
        if log.status == "dinas" or log.event_type == "DINAS":
            import json
            lampiran_path = "-"
            try:
                dinas_data = json.loads(log.device_id)
                keterangan = dinas_data.get("k", "Perizinan")
                shift_label = dinas_data.get("s", "Seharian")
                if dinas_data.get("f"):
                    base_url = str(request.base_url).rstrip("/")
                    lampiran_path = f"{base_url}/api/v1/uploads/permits/{dinas_data['f']}"
            except Exception:
                keterangan = log.device_id or "Izin Resmi"
                shift_label = "Seharian"
            
            group["in_time"] = "-"
            group["out_time"] = "-"
            group["status"] = "Izin"
            group["keterangan"] = keterangan
            group["device_id"] = "-"
            group["shift"] = shift_label
            group["lampiran"] = lampiran_path
        elif log.event_type == "IN":
            group["in_time"] = dt_wib.strftime("%H:%M:%S")
            group["shift"] = shift_label # Pastikan shift IN menjadi shift utama hari itu
            group["status"] = "Terlambat" if log.late else "Tepat Waktu"
            group["keterangan"] = "-"
            if group["device_id"] == "-":
                group["device_id"] = log.device_id or "-"
        elif log.event_type == "OUT":
            if group["out_time"] == "-":
                group["out_time"] = dt_wib.strftime("%H:%M:%S")
            if group["device_id"] == "-":
                group["device_id"] = log.device_id or "-"

    # Create Excel Workbook
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Laporan Kehadiran"

    headers = [
        "No", "Nama Pegawai", "NIY", "Tanggal", "Shift", "Jam Masuk", "Jam Pulang", "Status Kehadiran", "Keterangan", "Link Lampiran"
    ]
    ws.append(headers)

    # Style Header
    header_fill = PatternFill(start_color="3B82F6", end_color="3B82F6", fill_type="solid") # Tailwind blue-500
    header_font = Font(color="FFFFFF", bold=True)
    thin_border = Border(
        left=Side(style='thin'), right=Side(style='thin'), 
        top=Side(style='thin'), bottom=Side(style='thin')
    )
    center_aligned = Alignment(horizontal="center", vertical="center")

    for col in range(1, len(headers) + 1):
        cell = ws.cell(row=1, column=col)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = center_aligned
        cell.border = thin_border

    # Append Data
    for idx, g in enumerate(grouped_logs.values(), 1):
        user_name, employee_id = user_map.get(g["user_id"], ("Unknown", "-"))
        
        row_data = [
            idx,
            user_name,
            employee_id,
            g["date"],
            g["shift"],
            g["in_time"],
            g["out_time"],
            g["status"],
            g["keterangan"],
            g["lampiran"]
        ]
        ws.append(row_data)
        
        # Apply styles to the newly added row
        current_row = ws.max_row
        for col_idx in range(1, len(row_data) + 1):
            cell = ws.cell(row=current_row, column=col_idx)
            cell.border = thin_border
            
            # Center align specific columns: No(1), Tanggal(4), Shift(5), Jam Masuk(6), Jam Pulang(7)
            if col_idx in [1, 4, 5, 6, 7]:
                cell.alignment = center_aligned
                
            # Hyperlink for Lampiran (Column 10)
            if col_idx == 10 and cell.value != "-":
                cell.hyperlink = cell.value  # Set URL as hyperlink
                cell.value = "Lihat Bukti"
                cell.font = Font(color="0000EE", underline="single")
                cell.alignment = center_aligned

    # Auto-adjust column widths
    for col in ws.columns:
        max_length = 0
        column_letter = col[0].column_letter
        for cell in col:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        ws.column_dimensions[column_letter].width = max_length + 2

    # Save to BytesIO
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)

    # Generate filename with date range
    filename = "Laporan_Kehadiran_SISKA"
    if date_from:
        filename += f"_{date_from}"
    if date_to:
        filename += f"_sd_{date_to}"
    filename += ".xlsx"

    return StreamingResponse(
        iter([output.getvalue()]),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )

@router.get("/report")
async def get_attendance_report(
    date_from: str = Query(None, description="Start date (YYYY-MM-DD)"),
    date_to: str = Query(None, description="End date (YYYY-MM-DD)"),
    user_id: str = Query(None, description="Filter by user ID"),
    db: AsyncSession = Depends(get_db),
    _admin=Depends(get_current_admin),
):
    """Get attendance recap report in JSON format."""
    query = select(AttendanceLog).join(User).order_by(AttendanceLog.timestamp.desc())

    wib_tz = timezone(timedelta(hours=7))
    today = datetime.now(wib_tz).date()

    if not date_from:
        first_log = await db.execute(select(func.min(AttendanceLog.timestamp)))
        first_ts = first_log.scalar()
        if first_ts:
            date_from = first_ts.replace(tzinfo=timezone.utc).astimezone(wib_tz).strftime("%Y-%m-%d")
        else:
            date_from = today.replace(day=1).strftime("%Y-%m-%d")

    if not date_to:
        date_to = today.strftime("%Y-%m-%d")

    try:
        d_from = datetime.strptime(date_from, "%Y-%m-%d").date()
        start_dt = datetime.combine(d_from, datetime.min.time(), tzinfo=wib_tz).astimezone(timezone.utc)
        query = query.where(AttendanceLog.timestamp >= start_dt)
    except ValueError:
        raise HTTPException(status_code=400, detail="Format date_from harus YYYY-MM-DD")

    try:
        d_to = datetime.strptime(date_to, "%Y-%m-%d").date()
        end_dt = datetime.combine(d_to, datetime.max.time(), tzinfo=wib_tz).astimezone(timezone.utc)
        query = query.where(AttendanceLog.timestamp <= end_dt)
    except ValueError:
        raise HTTPException(status_code=400, detail="Format date_to harus YYYY-MM-DD")

    if user_id:
        query = query.where(AttendanceLog.user_id == user_id)

    # Base target users
    dept_order = case(
        (User.department.ilike('%Wadir%') | User.department.ilike('%Wakil Direktur%'), 2),
        (User.department.ilike('%Direktur%'), 1),
        (User.department.ilike('%Kaprodi%') | User.department.ilike('%Ketua Program%'), 3),
        (User.department.ilike('%Dosen%'), 4),
        (User.department.ilike('%Staf%') | User.department.ilike('%Staff%'), 5),
        else_=6
    )
    users_query = select(User.id, User.full_name, User.employee_id, User.department).where(User.is_active == True)
    if user_id:
        users_query = users_query.where(User.id == user_id)
        
    users_query = users_query.order_by(dept_order, User.department.asc(), User.full_name.asc())
        
    users_result = await db.execute(users_query)
    all_users = users_result.all()
    
    result = await db.execute(query)
    logs = result.scalars().all()

    # Calculate total working days in the range (cap at today to avoid future Alphas)
    calc_d_to = min(d_to, today)
    total_working_days = count_working_days(d_from, calc_d_to) if d_from <= calc_d_to else 0

    # Group logic
    grouped = {u.id: {"user_name": u.full_name, "employee_id": u.employee_id, "hadir": 0, "terlambat": 0, "izin": 0, "alpha": total_working_days, "logs_by_date": {}} for u in all_users}
    
    for log in logs:
        if not log.timestamp or log.user_id not in grouped:
            continue
        dt_wib = log.timestamp.replace(tzinfo=timezone.utc).astimezone(wib_tz)
        date_str = dt_wib.strftime("%Y-%m-%d")
        
        user_data = grouped[log.user_id]
        if date_str not in user_data["logs_by_date"]:
            user_data["logs_by_date"][date_str] = {"status": None, "in": None, "out": None, "late": False, "shift": None}
            
        daily = user_data["logs_by_date"][date_str]
        
        if log.status == "dinas" or log.event_type == "DINAS":
            daily["status"] = "Izin"
            import json
            try:
                dinas = json.loads(log.device_id)
                daily["shift"] = dinas.get("s", "Seharian")
            except:
                daily["shift"] = "Seharian"
        elif log.event_type == "IN":
            if daily["status"] != "Izin":
                daily["status"] = "Hadir"
            daily["in"] = dt_wib.strftime("%H:%M")
            if log.late:
                daily["late"] = True
                daily["status"] = "Terlambat"
            hour = dt_wib.hour
            daily["shift"] = "Shift Pagi" if hour < 15 else "Shift Sore"
        elif log.event_type == "OUT":
            daily["out"] = dt_wib.strftime("%H:%M")
            if not daily["shift"]:
                hour = dt_wib.hour
                daily["shift"] = "Shift Pagi" if hour < 15 else "Shift Sore"
                
    # Calculate recap
    recap = []
    for uid, data in grouped.items():
        hadir_count = 0
        terlambat_count = 0
        izin_count = 0
        
        detail_list = []
        for d, log_data in sorted(data["logs_by_date"].items()):
            if log_data["status"] == "Izin":
                izin_count += 1
            elif log_data["status"] == "Terlambat":
                hadir_count += 1
                terlambat_count += 1
            elif log_data["status"] == "Hadir":
                hadir_count += 1
                
            detail_list.append({
                "date": d,
                "in": log_data["in"] or "-",
                "out": log_data["out"] or "-",
                "shift": log_data["shift"] or "-",
                "status": log_data["status"] or "Hadir"
            })
            
        alpha_count = max(0, total_working_days - (hadir_count + izin_count))
            
        recap_item = {
            "user_id": uid,
            "user_name": data["user_name"],
            "employee_id": data["employee_id"],
            "hadir": hadir_count,
            "terlambat": terlambat_count,
            "izin": izin_count,
            "alpha": alpha_count,
            "total_hari_kerja": total_working_days
        }
        
        if user_id:
            recap_item["detail"] = detail_list
            
        recap.append(recap_item)

    # Maintain SQL query order
    
    def format_indo_date(date_str):
        months = ["Januari", "Februari", "Maret", "April", "Mei", "Juni", "Juli", "Agustus", "September", "Oktober", "November", "Desember"]
        try:
            y, m, d = date_str.split("-")
            return f"{int(d)} {months[int(m)-1]} {y}"
        except:
            return date_str

    period_str = f"{format_indo_date(date_from)} s.d. {format_indo_date(date_to)}"
    
    now_dt = datetime.now(wib_tz)
    months_list = ["Januari", "Februari", "Maret", "April", "Mei", "Juni", "Juli", "Agustus", "September", "Oktober", "November", "Desember"]
    print_date_indo = f"{now_dt.day} {months_list[now_dt.month - 1]} {now_dt.year}"

    return {
        "period": period_str,
        "print_date": print_date_indo,
        "is_single_user": user_id is not None,
        "recap": recap
    }

@router.get("/report/word")
async def export_attendance_word(
    date_from: str = Query(None, description="Start date (YYYY-MM-DD)"),
    date_to: str = Query(None, description="End date (YYYY-MM-DD)"),
    user_id: str = Query(None, description="Filter by user ID"),
    db: AsyncSession = Depends(get_db),
    _admin=Depends(get_current_admin),
):
    """Export recap report to Word (DOCX)."""
    # Use the same logic as /report to get data
    report_data = await get_attendance_report(date_from, date_to, user_id, db, _admin)
    
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.shared import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt
    import docx.shared
    import os

    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    
    # Set standard margins (1 inch all around)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.header_distance = Inches(0.4)
    
    # Move KOP SURAT to HEADER
    header = doc.sections[0].header
    
    # Clear default paragraph in header
    for p in header.paragraphs:
        p.text = ""
        p.paragraph_format.space_after = Pt(0)
    
    # Create invisible 3-column table in header
    table_kop = header.add_table(rows=1, cols=3, width=Inches(6.27))
    table_kop.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_kop.autofit = False
    
    # Set specific widths (Total ~6.27 inches for A4 with 1-inch margins)
    table_kop.columns[0].width = Inches(1.0)
    table_kop.columns[1].width = Inches(4.27)
    table_kop.columns[2].width = Inches(1.0)
    
    # Left logo
    cell_left = table_kop.cell(0, 0)
    cell_left.width = Inches(1.2)
    cell_left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_left = cell_left.paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_poltek = os.path.join(os.getcwd(), '..', 'frontend', 'public', 'logo-poltek.png')
    if os.path.exists(logo_poltek):
        p_left.add_run().add_picture(logo_poltek, width=Inches(1.0))
        
    # Middle text
    cell_mid = table_kop.cell(0, 1)
    cell_mid.width = Inches(5.1)
    cell_mid.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_mid = cell_mid.paragraphs[0]
    p_mid.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_mid.paragraph_format.space_after = Pt(0)
    
    run_y = p_mid.add_run("YAYASAN PENDIDIKAN BHAKTI PRAJA TEGAL\n")
    run_y.font.size = Pt(12)
    run_y.bold = True
    
    run_p = p_mid.add_run("POLITEKNIK BAJA TEGAL\n")
    run_p.font.size = Pt(14)
    run_p.font.color.rgb = docx.shared.RGBColor(0x00, 0x33, 0x66)
    run_p.bold = True
    
    run_a = p_mid.add_run('TERAKREDITASI "BAIK" NO. 341/SK/BAN-PT/Akred/PT/VII/2022\n')
    run_a.font.size = Pt(11)
    run_a.bold = True
    
    run2 = p_mid.add_run("Alamat : Jl. Raya Barat Dukuhwaru, Jatibarang-Slawi Km. 7, Kab. Tegal\nTelp. (0283) 6196309 - Website : www.pbjt.ac.id - E-mail : info@pbjt.ac.id")
    run2.font.size = Pt(9)
    
    # Right logo
    cell_right = table_kop.cell(0, 2)
    cell_right.width = Inches(1.0)
    cell_right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_right = cell_right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_banpt = os.path.join(os.getcwd(), '..', 'frontend', 'public', 'logo-banpt.png')
    if os.path.exists(logo_banpt):
        p_right.add_run().add_picture(logo_banpt, width=Inches(0.95))
        
    # Add bottom border 2 1/4 pt to Kop Table
    tbl = table_kop._tbl
    tblBorders = OxmlElement('w:tblBorders')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '18') # 18 eighths of a pt = 2.25 pt
    bottom.set(qn('w:space'), '0')
    bottom.set(qn('w:color'), 'auto')
    tblBorders.append(bottom)
    tbl.tblPr.append(tblBorders)
    
    # Empty paragraph for spacing below header
    header.add_paragraph()
    
    # MAIN DOCUMENT CONTENT
    
    # JUDUL
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("LAPORAN REKAPITULASI KEHADIRAN")
    run_title.bold = True
    run_title.underline = True
    run_title.font.size = Pt(14)
    
    doc.add_paragraph(f"Periode: {report_data['period']}")
    
    if report_data["is_single_user"] and len(report_data["recap"]) > 0:
        doc.add_paragraph(f"Nama: {report_data['recap'][0]['user_name']}")
        doc.add_paragraph(f"NIY: {report_data['recap'][0]['employee_id']}")
        
    doc.add_paragraph()
    
    # RECAP TABLE
    headers = ["No", "Nama Pegawai", "NIY", "Hadir", "Terlambat", "Izin", "Alpha"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Adjust widths to fit 6.27 inches perfectly
    widths = (Inches(0.4), Inches(2.27), Inches(1.2), Inches(0.6), Inches(0.6), Inches(0.6), Inches(0.6))
    
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].width = widths[i]
        p = hdr_cells[i].paragraphs[0]
        p.runs[0].bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    for idx, item in enumerate(report_data["recap"], 1):
        row_cells = table.add_row().cells
        for i in range(len(headers)):
            row_cells[i].width = widths[i]
            
        row_cells[0].text = str(idx)
        row_cells[1].text = item.get("user_name") or "-"
        row_cells[2].text = item.get("employee_id") or "-"
        row_cells[3].text = str(item.get("hadir", 0))
        row_cells[4].text = str(item.get("terlambat", 0))
        row_cells[5].text = str(item.get("izin", 0))
        row_cells[6].text = str(item.get("alpha", 0))
        
        # Center align numbers
        for col_idx in [0, 3, 4, 5, 6]:
            row_cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    doc.add_paragraph()
    
    # DETAIL TABLE FOR SINGLE USER
    if report_data["is_single_user"] and len(report_data["recap"]) > 0:
        detail = report_data["recap"][0].get("detail", [])
        if detail:
            doc.add_paragraph("Detail Kehadiran Harian:")
            d_headers = ["Tanggal", "Jam Masuk", "Jam Pulang", "Shift", "Status"]
            d_table = doc.add_table(rows=1, cols=len(d_headers))
            d_table.style = 'Table Grid'
            d_hdr_cells = d_table.rows[0].cells
            for i, h in enumerate(d_headers):
                d_hdr_cells[i].text = h
                d_hdr_cells[i].paragraphs[0].runs[0].bold = True
                
            for row in detail:
                cells = d_table.add_row().cells
                cells[0].text = row.get("date", "-")
                cells[1].text = row.get("in", "-")
                cells[2].text = row.get("out", "-")
                cells[3].text = row.get("shift", "-")
                cells[4].text = row.get("status", "-")
        doc.add_paragraph()
        
    # FOOTER (Tanda Tangan)
    from datetime import datetime
    
    now = datetime.now()
    bulan_indo = {
        1: "Januari", 2: "Februari", 3: "Maret", 4: "April",
        5: "Mei", 6: "Juni", 7: "Juli", 8: "Agustus",
        9: "September", 10: "Oktober", 11: "November", 12: "Desember"
    }
    tanggal_str = f"{now.day} {bulan_indo[now.month]} {now.year}"
    
    doc.add_paragraph()
    
    table_ttd = doc.add_table(rows=1, cols=2)
    table_ttd.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_ttd.autofit = False
    table_ttd.columns[0].width = Inches(4.0)
    table_ttd.columns[1].width = Inches(2.27)
    
    cell_sig = table_ttd.cell(0, 1)
    cell_sig.width = Inches(2.27)
    p_sig = cell_sig.paragraphs[0]
    p_sig.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sig.paragraph_format.space_after = Pt(0)
    
    run_date = p_sig.add_run(f"Dukuhwaru, {tanggal_str}\n")
    run_date.font.size = Pt(11)
    
    run_jab = p_sig.add_run("Wakil Direktur Bidang Akademik dan Kepegawaian\n\n\n\n\n")
    run_jab.font.size = Pt(11)
    
    run_nama = p_sig.add_run("Aziz Azindani, M.Kom\n")
    run_nama.bold = True
    run_nama.underline = True
    run_nama.font.size = Pt(11)
    
    run_nip = p_sig.add_run("850018701")
    run_nip.font.size = Pt(11)
    
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    
    filename = f"Laporan_Kehadiran_SISKA.docx"
    return StreamingResponse(
        iter([output.getvalue()]),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )

@router.post("/reset")
async def reset_attendance_logs(
    db: AsyncSession = Depends(get_db),
):
    """Reset attendance logs (DEBUG ONLY)"""
    from sqlalchemy import delete
    # Clear local logs
    await db.execute(delete(AttendanceLog))
    await db.commit()
    
    # Forward to AI API
    try:
        await ai_service.reset_attendance()
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"AI API error: {str(e)}")
        
    return {"status": "ok", "message": "Attendance logs reset"}
